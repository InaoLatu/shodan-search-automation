# Author: Inao Latourrette
# GitHub: https://github.com/InaoLatu
# LinkedIn: https://www.linkedin.com/in/inaolatourrette/
# Contact: inao.latourrette@gmail.com

# Description: It builds an automatic search in Shodan and a exploits discovery search
# The input strings corresponds the queries and exploits_query variables.
# It prints the results by console and creates a .docx with a table for each result of the search

from time import sleep
import shodan
from docx import Document

SHODAN_API_KEY = ""  # Insert your SHODAN_API_KEY
api = shodan.Shodan(SHODAN_API_KEY)
query = 'webcam 7  -401 http.component:"mootools"'  # query for automatic_search()
exploits_query = 'webcam 7'  # for exploits_search()
# searches parameters

def automatic_search():
    document = Document()

    document.add_heading('Shodan Search Automation', 0)

    document.add_paragraph(
        'This document contains the tables with the information related to each device found in the given searches in Shodan.'
        ' This whole document has been created by the python script used to do the automatic searches.  ')

    try:
        results = api.search(query=query)
        document.add_heading('Search: ' + query, level=1)
        document.add_paragraph("Tables for the search")

        for counter, result in enumerate(results['matches']):
            host = api.host(result['ip_str'])
            print("IP: {}\nOrganization: {}\nOperating System: {}".format(host['ip_str'], host.get('org', 'n/a'),
                                                                          host.get('os', 'n/a')))

            table_dict = {}
            table_dict["IP"] = host['ip_str']
            table_dict["Organization"] = host.get('org', 'n/a')

            location_string = ""
            if host['country_name'] is not None:
                location_string = location_string + host["country_name"]
            if host['city'] is not None:
                location_string = location_string + ", " + host["city"] + " "

            longitud_string = ""
            latitude_string = ""
            if host['longitude'] is not None:
                longitud_string = "Longitude: " + str(host['longitude']) + " "
            if host['latitude'] is not None:
                latitude_string = "Latitude: " + str(host['latitude']) + " "

            print(location_string)
            print(longitud_string)
            print(latitude_string)

            table_dict["Location"] = location_string
            table_dict['Longitude'] = longitud_string
            table_dict['Latitude'] = latitude_string

            ports_string = ""
            for item in host['data']:
                ports_string = ports_string + str(item['port']) + " "
            print("Ports: {}".format(ports_string))

            table_dict['Ports'] = ports_string

            data = host['data']
            products_string = ""
            for a in data:
                if "product" in a:
                    print("Product: {}".format(a['product']))
                    products_string = products_string + a['product'] + " "

            table_dict["Products"] = products_string
            print("")
            sleep(0.5)

            # Add the table for each result
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Key'
            hdr_cells[1].text = 'Value'

            for key, value in table_dict.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)

            document.add_paragraph("")
            if counter == 20:  # Number of devices per search
                break

        document.save('searches.docx')

    except shodan.APIError as e:
        print('Error {}'.format(e))


def exploits_search():
    try:
        exploits = api.exploits.search(query=exploits_query)
        document = Document()

        document.add_heading('Shodan Exploits Search Automation', 0)
        document.add_heading('Search: ' + exploits_query, level=1)

        for e in exploits['matches']:
            exploit_string = ""
            if e['source'] == "CVE":
                exploit_string = str(e['source']) + " - " + str(e['cve'])
            else:
                exploit_string = str(e['source'])

            exploit_string = exploit_string + ": " + e['description']
            print(exploit_string)

            document.add_paragraph(exploit_string)

        document.save("exploits.docx")

    except shodan.APIError as e:
        print('Error {}'.format(e))


def main():
    automatic_search()
    exploits_search()


if __name__ == '__main__':
    main()
